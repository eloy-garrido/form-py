#!/usr/bin/env python3
"""
MCP Server para procesamiento de plantillas Word - Versión dinámica
Autodetecta y reemplaza TODOS los campos {{campo}} en documentos Word
"""

import os
import re
import json
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List

from docx import Document
from mcp.server.fastmcp import FastMCP

# Configuración de paths
BASE_DIR = Path(__file__).parent
TEMPLATE_PATH = BASE_DIR / "plantilla_formulario.docx"
OUTPUT_DIR = BASE_DIR / "form-generados"

# Crear servidor MCP
mcp = FastMCP("Word Form Processor - Dynamic Fields")

def find_template_fields(doc_path: Path) -> List[str]:
    """
    Extrae TODOS los campos únicos {{campo}} de un documento Word
    """
    try:
        doc = Document(doc_path)
        fields = set()
        
        def extract_fields(text: str):
            # Encuentra todos los patrones {{campo}} incluyendo espacios alrededor del nombre
            matches = re.findall(r'\{\{\s*([^{}]+?)\s*\}\}', text)
            fields.update(matches)
        
        # Buscar en todos los párrafos
        for paragraph in doc.paragraphs:
            extract_fields(paragraph.text)
        
        # Buscar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        extract_fields(paragraph.text)
        
        # Buscar en headers y footers
        for section in doc.sections:
            for part in [section.header, section.footer]:
                if part:
                    for paragraph in part.paragraphs:
                        extract_fields(paragraph.text)
        
        return sorted(list(fields))
    
    except Exception as e:
        raise ValueError(f"Error al leer la plantilla: {str(e)}")

def replace_fields_in_document(doc: Document, fields_data: Dict[str, str]) -> int:
    """
    Reemplaza dinámicamente TODOS los campos encontrados en el documento
    Retorna el número de reemplazos realizados
    """
    replacements_count = 0
    
    def replace_in_text(text: str) -> (str, int):
        """Reemplaza campos en un texto y devuelve el texto modificado y el conteo"""
        count = 0
        new_text = text
        
        # Encuentra todos los campos en el texto actual
        fields_in_text = re.findall(r'\{\{\s*([^{}]+?)\s*\}\}', text)
        
        for field in fields_in_text:
            clean_field = field.strip()  # Elimina espacios alrededor del nombre del campo
            if clean_field in fields_data:
                # Mantiene el formato original del campo (con/sin espacios)
                original_pattern = f"{{{{ {field} }}}}" if ' ' in field else f"{{{{{field}}}}}"
                new_text = new_text.replace(original_pattern, str(fields_data[clean_field]))
                count += 1
        
        return new_text, count
    
    # Reemplazar en párrafos principales
    for paragraph in doc.paragraphs:
        new_text, count = replace_in_text(paragraph.text)
        if count > 0:
            paragraph.text = new_text
            replacements_count += count
    
    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    new_text, count = replace_in_text(paragraph.text)
                    if count > 0:
                        paragraph.text = new_text
                        replacements_count += count
    
    # Reemplazar en headers y footers
    for section in doc.sections:
        for part in [section.header, section.footer]:
            if part:
                for paragraph in part.paragraphs:
                    new_text, count = replace_in_text(paragraph.text)
                    if count > 0:
                        paragraph.text = new_text
                        replacements_count += count
    
    return replacements_count

@mcp.tool()
def list_template_fields() -> str:
    """
    Lista TODOS los campos {{campo}} detectados dinámicamente en la plantilla
    """
    if not TEMPLATE_PATH.exists():
        return f"❌ Error: No se encontró la plantilla en {TEMPLATE_PATH}"
    
    try:
        fields = find_template_fields(TEMPLATE_PATH)
        
        if not fields:
            return "ℹ️ No se encontraron campos {{}} en la plantilla"
        
        result = f"📝 Campos detectados en la plantilla ({len(fields)}):\n\n"
        for i, field in enumerate(fields, 1):
            result += f"{i}. {{{{{field}}}}}\n"
        
        return result
    
    except Exception as e:
        return f"❌ Error: {str(e)}"

@mcp.tool()
def generate_document(fields_data: str, filename_prefix: str = "") -> str:
    """
    Genera documento reemplazando TODOS los campos detectados automáticamente
    
    Args:
        fields_data: JSON string con los datos para los campos detectados
                    Ejemplo: '{"nombre": "Juan Pérez", "rut": "12345678-9"}'
        filename_prefix: Prefijo opcional para el nombre del archivo
    """
    try:
        data = json.loads(fields_data)
        
        if not TEMPLATE_PATH.exists():
            return f"❌ Error: No se encontró la plantilla en {TEMPLATE_PATH}"
        
        OUTPUT_DIR.mkdir(exist_ok=True)
        doc = Document(TEMPLATE_PATH)
        
        replacements_count = replace_fields_in_document(doc, data)
        
        if replacements_count == 0:
            return "⚠️ Advertencia: No se realizaron reemplazos. Verifica que los nombres coincidan con los campos detectados."
        
        # Generar nombre de archivo único
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{filename_prefix}_{timestamp}.docx" if filename_prefix else f"documento_{timestamp}.docx"
        output_path = OUTPUT_DIR / filename
        doc.save(output_path)
        
        # Obtener lista de campos reemplazados
        replaced_fields = [field for field in find_template_fields(TEMPLATE_PATH) if field in data]
        
        result = f"✅ Documento generado exitosamente!\n\n"
        result += f"📁 Archivo: {filename}\n"
        result += f"📍 Ubicación: {output_path}\n"
        result += f"🔄 Campos reemplazados: {replacements_count}\n\n"
        result += f"📝 Campos detectados y reemplazados:\n"
        for field in replaced_fields:
            result += f"   • {field}: {data[field]}\n"
        
        return result
    
    except json.JSONDecodeError:
        return "❌ Error: Los datos deben estar en formato JSON válido"
    except Exception as e:
        return f"❌ Error: {str(e)}"

@mcp.tool()
def preview_replacements(fields_data: str) -> str:
    """
    Muestra cómo quedarían los reemplazos para los campos detectados
    """
    try:
        data = json.loads(fields_data)
        
        if not TEMPLATE_PATH.exists():
            return f"❌ Error: No se encontró la plantilla en {TEMPLATE_PATH}"
        
        template_fields = find_template_fields(TEMPLATE_PATH)
        result = "🔍 Vista previa de reemplazos:\n\n"
        
        replaced_count = 0
        for field in template_fields:
            if field in data:
                result += f"✅ {{{{{field}}}}} → {data[field]}\n"
                replaced_count += 1
            else:
                result += f"⚠️ {{{{{field}}}}} → [SIN VALOR]\n"
        
        extra_fields = set(data.keys()) - set(template_fields)
        if extra_fields:
            result += f"\n📝 Campos en los datos no encontrados en la plantilla:\n"
            for field in extra_fields:
                result += f"ℹ️ {field}: {data[field]}\n"
        
        result += f"\n📊 Resumen: {replaced_count}/{len(template_fields)} campos serían reemplazados"
        return result
    
    except json.JSONDecodeError:
        return "❌ Error: Los datos deben estar en formato JSON válido"
    except Exception as e:
        return f"❌ Error: {str(e)}"

if __name__ == "__main__":
    mcp.run()