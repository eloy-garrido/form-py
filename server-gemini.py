#!/usr/bin/env python3
"""
MCP Server para procesamiento de plantillas Word - Versión que funciona con Claude
Reemplaza campos {{campo}} en documentos Word con datos proporcionados
"""

import os
import re
import json
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
# Es buena práctica importar el tipo Document directamente desde docx
from docx.document import Document as DocxDocument
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from mcp.server.fastmcp import FastMCP

# Configuración de paths
BASE_DIR = Path(__file__).parent
TEMPLATE_PATH = BASE_DIR / "plantilla_formulario.docx"
OUTPUT_DIR = BASE_DIR / "form-generados"

# Crear servidor MCP
mcp = FastMCP("Word Form Processor - Claude Compatible")


def find_template_fields(doc_path: Path) -> List[str]:
    """
    Extrae todos los campos {{campo}} de un documento Word de manera eficiente
    """
    try:
        doc = Document(doc_path)
        fields = set()
        
        def extract_fields_from_text(text: str):
            """Extrae campos de un texto dado"""
            matches = re.findall(r'\{\{([^}]+)\}\}', text)
            fields.update(matches)
        
        # Buscar en párrafos principales
        for paragraph in doc.paragraphs:
            extract_fields_from_text(paragraph.text)
        
        # Buscar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Las celdas pueden contener múltiples párrafos
                    for paragraph in cell.paragraphs:
                        extract_fields_from_text(paragraph.text)
        
        # Buscar en headers y footers
        for section in doc.sections:
            # Header principal (primera página)
            if section.header:
                for paragraph in section.header.paragraphs:
                    extract_fields_from_text(paragraph.text)
            
            # Header para páginas impares/pares (si se usan)
            if section.even_page_header:
                for paragraph in section.even_page_header.paragraphs:
                    extract_fields_from_text(paragraph.text)

            if section.first_page_header:
                for paragraph in section.first_page_header.paragraphs:
                    extract_fields_from_text(paragraph.text)
            
            # Footer principal (primera página)
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    extract_fields_from_text(paragraph.text)

            # Footer para páginas impares/pares (si se usan)
            if section.even_page_footer:
                for paragraph in section.even_page_footer.paragraphs:
                    extract_fields_from_text(paragraph.text)

            if section.first_page_footer:
                for paragraph in section.first_page_footer.paragraphs:
                    extract_fields_from_text(paragraph.text)
        
        return sorted(list(fields))
    
    except Exception as e:
        raise ValueError(f"Error al leer la plantilla: {str(e)}")


def replace_fields_in_document(doc: DocxDocument, fields_data: Dict[str, str]) -> int:
    """
    Reemplaza todos los campos en el documento de manera eficiente,
    manejando campos que pueden estar divididos en múltiples 'runs'.
    Retorna el número de reemplazos realizados.
    """
    replacements_count = 0
    
    def _replace_in_paragraph_or_cell(paragraph: DocxParagraph, data_map: Dict[str, str]):
        nonlocal replacements_count
        original_text = paragraph.text
        new_text = original_text
        current_paragraph_replacements = 0

        # Iterar sobre los datos para reemplazar todos los campos en este texto
        for field, value in data_map.items():
            # Usar re.escape para el nombre del campo, por si contiene caracteres especiales de regex
            pattern = r'\{\{' + re.escape(field) + r'\}\}'
            
            # Contar las ocurrencias antes de reemplazar
            matches = list(re.finditer(pattern, new_text))
            if matches:
                new_text = re.sub(pattern, str(value), new_text)
                current_paragraph_replacements += len(matches)
        
        # Si se realizó al menos un reemplazo, actualiza el párrafo
        # Asignar a paragraph.text automáticamente gestiona los 'runs' internos.
        if current_paragraph_replacements > 0:
            paragraph.text = new_text
            replacements_count += current_paragraph_replacements

    # Reemplazar en párrafos principales
    for paragraph in doc.paragraphs:
        _replace_in_paragraph_or_cell(paragraph, fields_data)
    
    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph_or_cell(paragraph, fields_data)
    
    # Reemplazar en headers y footers para todas las secciones
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                _replace_in_paragraph_or_cell(paragraph, fields_data)
        
        if section.even_page_header:
            for paragraph in section.even_page_header.paragraphs:
                _replace_in_paragraph_or_cell(paragraph, fields_data)

        if section.first_page_header:
            for paragraph in section.first_page_header.paragraphs:
                _replace_in_paragraph_or_cell(paragraph, fields_data)
        
        if section.footer:
            for paragraph in section.footer.paragraphs:
                _replace_in_paragraph_or_cell(paragraph, fields_data)

        if section.even_page_footer:
            for paragraph in section.even_page_footer.paragraphs:
                _replace_in_paragraph_or_cell(paragraph, fields_data)

        if section.first_page_footer:
            for paragraph in section.first_page_footer.paragraphs:
                _replace_in_paragraph_or_cell(paragraph, fields_data)
    
    return replacements_count


@mcp.tool()
def list_template_fields() -> str:
    """
    Lista todos los campos {{campo}} disponibles en la plantilla
    """
    if not TEMPLATE_PATH.exists():
        return f"❌ Error: No se encontró la plantilla en {TEMPLATE_PATH}"
    
    try:
        fields = find_template_fields(TEMPLATE_PATH)
        
        if not fields:
            return "ℹ️ No se encontraron campos {{}} en la plantilla"
        
        result = f"📝 Campos encontrados en la plantilla ({len(fields)}):\n\n"
        for i, field in enumerate(fields, 1):
            result += f"{i}. {{{{{field}}}}}\n"
        
        return result
    
    except Exception as e:
        return f"❌ Error: {str(e)}"


@mcp.tool()
def preview_replacements(fields_data: str) -> str:
    """
    Previsualiza qué campos se reemplazarían sin generar el archivo
    
    Args:
        fields_data: JSON string con los datos de los campos
                    Ejemplo: '{"nombre_estudiante": "Juan Pérez", "rut_estudiante": "12345678-9"}'
    """
    try:
        data = json.loads(fields_data)
        
        if not TEMPLATE_PATH.exists():
            return f"❌ Error: No se encontró la plantilla en {TEMPLATE_PATH}"
        
        # Obtener campos de la plantilla
        template_fields = find_template_fields(TEMPLATE_PATH)
        
        result = "🔍 Vista previa de reemplazos:\n\n"
        
        # Mostrar campos que se reemplazarían
        replaced_fields = []
        for field in template_fields:
            if field in data:
                result += f"✅ {{{{{field}}}}} → '{data[field]}'\n" # Añadidas comillas para claridad
                replaced_fields.append(field)
            else:
                result += f"⚠️ {{{{{field}}}}} → [SIN VALOR]\n"
        
        # Mostrar campos extra en los datos
        extra_fields = set(data.keys()) - set(template_fields)
        if extra_fields:
            result += f"\n📝 Campos en los datos que no están en la plantilla:\n"
            for field in extra_fields:
                result += f"ℹ️ {field}: '{data[field]}'\n"
        
        result += f"\n📊 Resumen: {len(replaced_fields)}/{len(template_fields)} campos se reemplazarían"
        
        return result
    
    except json.JSONDecodeError:
        return "❌ Error: Los datos deben estar en formato JSON válido. Asegúrate de que es una cadena JSON."
    except Exception as e:
        return f"❌ Error: {str(e)}"


@mcp.tool()
def generate_document(fields_data: str, filename_prefix: str = "") -> str:
    """
    Genera un nuevo documento Word reemplazando los campos con los datos proporcionados
    
    Args:
        fields_data: JSON string con los datos de los campos
                    Ejemplo: '{"nombre_estudiante": "Juan Pérez", "rut_estudiante": "12345678-9"}'
        filename_prefix: Prefijo opcional para el nombre del archivo (default: "")
    """
    try:
        data = json.loads(fields_data)
        
        if not TEMPLATE_PATH.exists():
            return f"❌ Error: No se encontró la plantilla en {TEMPLATE_PATH}"
        
        # Asegurar que existe el directorio de salida
        OUTPUT_DIR.mkdir(exist_ok=True)
        
        # Cargar documento
        doc = Document(TEMPLATE_PATH)
        
        # Realizar reemplazos
        replacements_count = replace_fields_in_document(doc, data)
        
        # Generar nombre de archivo
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if filename_prefix:
            # Asegurar que no hay guion bajo duplicado y limpiar el prefijo
            clean_prefix = re.sub(r'[^\w\s-]', '', filename_prefix).strip()
            clean_prefix = re.sub(r'[-\s]+', '_', clean_prefix)
            if clean_prefix: # Evitar doble guion bajo si el prefijo está vacío
                filename = f"{clean_prefix}_{timestamp}.docx"
            else:
                filename = f"documento_{timestamp}.docx"
        else:
            # Intentar usar nombre_estudiante como prefijo si está disponible
            name = data.get('nombre_estudiante', data.get('nombre', ''))
            if name:
                # Limpiar nombre para uso en archivo
                clean_name = re.sub(r'[^\w\s-]', '', name).strip()
                clean_name = re.sub(r'[-\s]+', '_', clean_name)
                filename = f"{clean_name}_{timestamp}.docx"
            else:
                filename = f"documento_{timestamp}.docx"
        
        # Guardar archivo
        output_path = OUTPUT_DIR / filename
        doc.save(output_path)
        
        result = f"✅ Documento generado exitosamente!\n\n"
        result += f"📁 Archivo: {filename}\n"
        result += f"📍 Ubicación: {output_path}\n"
        result += f"🔄 Reemplazos realizados: {replacements_count}\n"
        result += f"📊 Campos procesados: {len(data)}\n\n"
        result += f"🎯 Datos procesados:\n"
        # Limitar la longitud de los valores para no saturar el output si son muy largos
        for key, value in data.items():
            display_value = str(value)
            if len(display_value) > 100:
                display_value = display_value[:97] + "..."
            result += f"   • {key}: '{display_value}'\n"
        
        return result
    
    except json.JSONDecodeError:
        return "❌ Error: Los datos deben estar en formato JSON válido. Asegúrate de que es una cadena JSON."
    except Exception as e:
        return f"❌ Error: {str(e)}"


@mcp.tool()
def create_pascale_document() -> str:
    """
    Función especializada para crear el documento de Pascale con datos predefinidos
    """
    try:
        # Datos predefinidos de Pascale
        data = {
            "nombre_estudiante": "Pascale Cataleya Figueroa Olguín",
            "rut_estudiante": "26.292.242-1",
            "nacimiento_estudiante": "27-05-2018",
            "edad_estudiante_años_meses": "7 años 0 meses",
            "curso": "1°A"
        }
        
        if not TEMPLATE_PATH.exists():
            return f"❌ Error: No se encontró la plantilla en {TEMPLATE_PATH}"
        
        # Asegurar que existe el directorio de salida
        OUTPUT_DIR.mkdir(exist_ok=True)
        
        # Cargar documento
        doc = Document(TEMPLATE_PATH)
        
        # Realizar reemplazos
        replacements_count = replace_fields_in_document(doc, data)
        
        # Generar nombre de archivo
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Informe_Pascale_{timestamp}.docx"
        
        # Guardar archivo
        output_path = OUTPUT_DIR / filename
        doc.save(output_path)
        
        result = f"✅ Documento de Pascale generado exitosamente!\n\n"
        result += f"📁 Archivo: {filename}\n"
        result += f"📍 Ubicación: {output_path}\n"
        result += f"🔄 Reemplazos realizados: {replacements_count}\n"
        result += f"📊 Campos procesados: {len(data)}\n\n"
        result += f"🎯 Datos utilizados:\n"
        for key, value in data.items():
            result += f"   • {key}: '{value}'\n"
        
        return result
    
    except Exception as e:
        return f"❌ Error: {str(e)}"


@mcp.tool()
def debug_info() -> str:
    """
    Información de debug del sistema
    """
    result = "🔧 Información de Debug\n\n"
    result += f"📁 Directorio base: {BASE_DIR}\n"
    result += f"📄 Plantilla: {TEMPLATE_PATH}\n"
    result += f"✅ Plantilla existe: {TEMPLATE_PATH.exists()}\n"
    if TEMPLATE_PATH.exists():
        try:
            # Intentar abrir la plantilla para ver si es un archivo docx válido
            Document(TEMPLATE_PATH)
            result += "✅ Plantilla es un archivo DOCX válido\n"
        except Exception as e:
            result += f"❌ Error al abrir plantilla DOCX: {e}\n"

    result += f"📁 Directorio salida: {OUTPUT_DIR}\n"
    result += f"✅ Directorio salida existe: {OUTPUT_DIR.exists()}\n\n"
    
    if TEMPLATE_PATH.exists():
        try:
            fields = find_template_fields(TEMPLATE_PATH)
            result += f"🔍 Campos en plantilla ({len(fields)}):\n"
            if fields:
                for field in fields:
                    result += f"   • {{{{{field}}}}}\n"
            else:
                result += "   (Ninguno encontrado)\n"
        except Exception as e:
            result += f"❌ Error leyendo campos de plantilla: {e}\n"
    
    result += f"\n🐍 Versión Python: {sys.version.split()[0]}\n"
    result += f"📦 Directorio actual: {os.getcwd()}\n"
    
    # Verificar dependencias
    try:
        import docx
        result += f"✅ python-docx: Disponible (Versión: {docx.__version__ if hasattr(docx, '__version__') else 'Desconocida'})\n"
    except ImportError:
        result += f"❌ python-docx: No disponible\n"
    
    try:
        import mcp
        result += f"✅ mcp: Disponible (Versión: {mcp.__version__ if hasattr(mcp, '__version__') else 'Desconocida'})\n"
    except ImportError:
        result += f"❌ mcp: No disponible\n"
    
    return result


@mcp.tool()
def list_generated_documents() -> str:
    """
    Lista todos los documentos generados en el directorio de salida
    """
    try:
        if not OUTPUT_DIR.exists():
            return "📁 El directorio de documentos generados aún no existe"
        
        # Buscar archivos .docx
        docx_files = list(OUTPUT_DIR.glob("*.docx"))
        
        if not docx_files:
            return "📁 No hay documentos generados aún"
        
        # Ordenar por fecha de modificación (más reciente primero)
        docx_files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
        
        result = f"📋 Documentos generados ({len(docx_files)}):\n\n"
        
        for i, file_path in enumerate(docx_files, 1):
            # Obtener información del archivo
            stat = file_path.stat()
            size_mb = stat.st_size / (1024 * 1024)
            modified_time = datetime.fromtimestamp(stat.st_mtime)
            
            result += f"{i}. {file_path.name}\n"
            result += f"   📅 Modificado: {modified_time.strftime('%Y-%m-%d %H:%M:%S')}\n"
            result += f"   📊 Tamaño: {size_mb:.2f} MB\n\n"
        
        return result
    
    except Exception as e:
        return f"❌ Error: {str(e)}"


if __name__ == "__main__":
    # Asegúrate de que el directorio de salida exista al iniciar el servidor
    # o al menos antes de la primera generación.
    # OUTPUT_DIR.mkdir(exist_ok=True) # Ya se hace en generate_document y create_pascale_document

    # Ejecutar el servidor
    mcp.run()